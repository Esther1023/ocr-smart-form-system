from docx import Document
import os

def modify_template(template_path: str, output_path: str):
    """
    在Word模板中添加变量标记
    """
    doc = Document(template_path)
    
    # 定义需要标记的文本及其对应的变量
    replacements = {
        # 甲方信息
        "甲方：": "甲方：{{ company_name }}",
        "税号：": "税号：{{ tax_number }}",
        "注册地址：": "注册地址：{{ reg_address }}",
        "注册电话：": "注册电话：{{ reg_phone }}",
        "开户行：": "开户行：{{ bank_name }}",
        "账号：": "账号：{{ bank_account }}",
        "联系人：": "联系人：{{ contact_name }}",
        "联系电话：": "联系电话：{{ contact_phone }}",
        "邮寄地址：": "邮寄地址：{{ mail_address }}",
        "简道云账号：": "简道云账号：{{ jdy_account }}",
        
        # 服务期限
        "1      年": "{{ service_years }} 年",
        "2024   年": "{{ start_year }} 年",
        "月     日": "{{ start_month }} 月 {{ start_day }} 日",
        "2025   年": "{{ end_year }} 年",
        "月     日": "{{ end_month }} 月 {{ end_day }} 日",
        
        # 费用信息
        "XXXX": "{{ total_amount }}",
        "壹": "{{ total_amount_cn }}",
        "6": "{{ tax_rate }}",
        "XXX": "{{ payment_amount }}",
        "XXXX": "{{ payment_amount_cn }}",
        "365元/人/年": "{{ unit_price }}元/人/年",
        
        # 其他信息
        "XX人": "{{ user_count }}人",
        "XX": "{{ payment_days }}",
        "普通/专用": "{{ invoice_type }}"
    }
    
    # 遍历所有段落
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
    
    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
    
    # 保存修改后的模板
    doc.save(output_path)
    print(f"模板已修改并保存为：{output_path}")

if __name__ == "__main__":
    # 获取当前目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 设置输入和输出路径
    template_path = os.path.join(current_dir, "合同模板.docx")
    output_path = os.path.join(current_dir, "合同模板_带变量.docx")
    
    # 修改模板
    modify_template(template_path, output_path)
